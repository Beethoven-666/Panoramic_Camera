from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\Panoramic_Camera\outputs\Gemini305_RTABMap_Windows_Photo_and_Video_Project_Proposal.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT = "F4F6F9"
HEADER = "E8EEF5"
GRAY = "666666"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def add_bottom_border(paragraph, color=BLUE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_text(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for cell, header in zip(header_row.cells, headers):
        set_cell_shading(cell, HEADER)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(body)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(10.5)
        s.paragraph_format.line_spacing = 1.2
        s.paragraph_format.space_after = Pt(4)


def setup_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section):
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Gemini 305 RGB-D 三维建图与侧扫全景项目书")
    set_run_font(r, size=8.5, color=GRAY)
    add_bottom_border(p, color="D7DBE2", size="4")
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=8.5, color=GRAY)
    add_page_field(p)
    r = p.add_run(" 页")
    set_run_font(r, size=8.5, color=GRAY)


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    set_run_font(r, size=16, color=BLUE, bold=True)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    set_run_font(r, size=13, color=BLUE, bold=True)
    return p


def h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    set_run_font(r, size=11.5, color=DARK_BLUE, bold=True)
    return p


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("项目实施书")
    set_run_font(r, size=15, color=GRAY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Gemini 305 RGB-D 三维建图与\n移动侧扫全景系统")
    set_run_font(r, size=27, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Windows 采集、RTAB-Map 三维建图与深度约束二维展开方案")
    set_run_font(r, size=13, color=DARK_BLUE)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    values = [
        ("项目代号", "G305-RTABMAP-PANO"),
        ("实施平台", "Windows 10/11 原生运行环境（采集、建图、渲染与交付均在同一 Windows 主机完成）"),
        ("项目状态", "建设方案：覆盖 RGB-D 采集、三维建图、二维侧扫全景、质量控制与交付"),
        ("编制日期", "2026 年 7 月 13 日"),
    ]
    for i, (label, value) in enumerate(values):
        set_cell_shading(table.cell(i, 0), HEADER)
        p = table.cell(i, 0).paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, bold=True, color=INK)
        p = table.cell(i, 1).paragraphs[0]
        r = p.add_run(value)
        set_run_font(r)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用范围：室内或半室内、场景基本静止、相机连续单向水平侧移的 RGB-D 全景采集")
    set_run_font(r, size=10, color=GRAY, italic=True)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    setup_section(section)
    setup_styles(doc)
    add_header_footer(section)
    add_cover(doc)

    h1(doc, "一、执行摘要")
    add_text(doc, "本项目面向奥比中光 Gemini 305 RGB-D 相机，在 Windows 环境下构建一条可审计、可失败关闭（fail-closed）的移动侧扫全景生产链路。系统先对同步、标定且对齐到彩色坐标系的 RGB-D 序列进行三维位姿估计和图优化，再将原始全分辨率 RGB-D 依据优化后的 SE(3) 位姿重投影到统一的正射侧扫平面，最后通过深度约束的 GraphCut 接缝、严格单 owner 分区和局部 MultiBand 融合生成二维全景图。")
    add_text(doc, "本项目建设范围包括 Windows RGB-D 采集、严格会话契约、RGB-D 位姿估计与图优化、RTAB-Map 三维建图、深度正射投影、接缝融合、质量门禁和原子交付。所有模块均按本项目实施计划开发、联调和验收；RTAB-Map 负责三维建图与优化位姿，二维全景由原始 RGB-D 的正射重投影生成。")
    add_callout(doc, "核心决策", "二维成果不能通过对三维网格截图或简单纹理展开获得。正式全景必须使用原始 RGB-D + RTAB-Map 优化位姿进行正射重投影，才能正确表达近景视差、遮挡与深度断层。")

    h2(doc, "项目目标")
    add_bullets(doc, [
        "在 Windows 主机复用 Flash 的设备枚举、profile 选择、D2C 对齐、帧同步及照片触发能力，提供原生连续视频与无界面照片采集两种 RGB-D 会话采集方式。",
        "通过 RTAB-Map 生成可复核的三维点云、可选网格和优化后的关键帧位姿图。",
        "将经过审计的 camera_to_world 4×4 SE(3) 位姿输入现有正射侧扫渲染器，输出无遮挡拼贴、无黑边和无几何伪造的二维全景。",
        "在位姿、深度、接缝、覆盖或资源预算不满足要求时明确失败，不发布部分或猜测性结果。",
    ])

    h2(doc, "范围与非范围")
    add_table(doc, ["范围", "说明"], [
        ("正式范围", "室内或半室内侧向移动扫描；场景基本静止；最近物体约 0.5 m；速度目标不高于 1.5 m/s。"),
        ("输入范围", "Gemini 305 对齐 RGB、对齐深度、相机内参、深度比例、时间戳和曝光元数据构成的严格会话。"),
        ("输出范围", "RTAB-Map 数据库、点云/可选网格、优化位姿、正射侧扫全景、质量报告和交付清单。"),
        ("非范围", "单目/RGB-only 拼接、二维单应矩阵拼接、时间插值伪造位姿、GPS 导航、室外大尺度测绘、动态场景重建。"),
    ], [2100, 7260])

    h1(doc, "二、建设内容与功能范围")
    add_text(doc, "本章按完整建设范围描述系统能力，不以任何模块的当前代码状态作为项目边界。每一工作包均须完成设计、开发、测试、联调、异常处理与验收，最终形成可在 Windows 环境重复运行的正式交付链路。")
    add_table(doc, ["工作包", "建设内容", "验收产物"], [
        ("RGB-D 采集", "Gemini 305 同步采集、软件 D2C 对齐、曝光控制、异步写盘、会话清理与元数据记录。", "会话目录、frames.csv、calibration.json、采集报告。"),
        ("原生视频采集", "复用 Flash 中已验证的设备枚举、profile 选择、D2C 对齐与帧同步能力，启动连续 RGB-D 视频流并逐帧写入正式会话。", "原生视频会话、帧同步/丢帧/曝光记录。"),
        ("无界面照片采集", "复用 Flash 的软件触发、同步取帧和单帧写盘能力；以 CLI/API 执行单次采集，不实现桌面入口。", "照片 RGB-D 会话、每次触发的完整元数据与失败记录。"),
        ("会话契约", "校验 aligned depth、内参、深度单位、尺寸、路径、时间戳和曝光；禁止 raw depth 冒充对齐深度。", "结构校验报告、失败样例测试。"),
        ("质量分析", "评估曝光、清晰度、纹理、主扫描段、相邻运动和渲染源布局。", "质量评分、扫描段与源帧选择报告。"),
        ("位姿图", "实现 RGB-D odometry、相邻/有限非相邻 RGB-D 边、信息矩阵、SE(3) 与轨迹质量审计。", "transforms.json、位姿图质量报告。"),
        ("RTAB-Map 三维建图", "实现 Windows 原生 RTAB-Map 桥接、数据库导出、关键帧关联、点云/网格与优化位姿适配。", "map.db、cloud.ply、可选 mesh、审计 JSON。"),
        ("二维投影与融合", "实现原始全分辨率 RGB-D 正射重投影、深度硬约束 owner、GraphCut 与局部 MultiBand。", "panorama.jpg、render_transforms.json、接缝审计。"),
        ("交付与测试", "实现 pending + 原子发布、失败报告、单元/集成/现场验收测试。", "report.json、delivery.json、测试与验收报告。"),
    ], [1800, 5050, 2510])

    h2(doc, "目标正式处理链路")
    add_numbers(doc, [
        "采集或接收正式 RGB-D 会话；验证标定、对齐、深度比例、尺寸、曝光与时间戳。",
        "评估输入画质与扫描运动；从序列中选择具有真实覆盖的 RGB-D pose nodes。",
        "运行 Open3D RGB-D odometry 并优化仅由 RGB-D 边构成的 pose graph。",
        "把选定原始全分辨率 RGB-D 根据优化后的位姿正射重投影为侧扫条带。",
        "在相邻重叠走廊中执行深度约束 GraphCut、owner 拓扑审计和局部 MultiBand 融合。",
        "裁剪、复核质量门禁，并以 delivery.json 最后发布的方式完成交付。",
    ])

    h1(doc, "三、总体技术方案")
    h2(doc, "总体架构")
    add_text(doc, "系统采用全 Windows 原生架构。采集模块以无界面库的形式复用 Flash 中已经验证的设备控制能力，提供 Gemini 305 原生连续视频流与软件触发照片两种正式输入模式；两种模式写入相同的严格 RGB-D 会话格式。建图模块以 Windows 原生 C++ RTAB-Map 桥接程序处理会话目录，渲染与交付继续由 Windows Python 程序完成。所有组件在同一 Windows 主机内安装、运行和验收。")
    add_table(doc, ["层级", "运行位置", "职责", "关键产物"], [
        ("采集层", "Windows 原生 Python + Flash 功能复用", "使用 Flash 的设备枚举、profile、D2C 对齐、帧同步和软件触发能力，支持视频与无界面照片两种 RGB-D 采集。", "统一会话目录、frames.csv、calibration.json"),
        ("建图层", "Windows 原生 C++ + RTAB-Map", "读取会话，运行 RGB-D 里程计、图优化、回环和三维地图导出。", "map.db、点云、网格、优化位姿"),
        ("适配层", "Windows 原生 C++/Python", "把 RTAB-Map 关键帧与原始 frame_id 精确关联，导出 mm 单位 SE(3) 和审计数据。", "rtabmap_transforms.json"),
        ("渲染层", "Windows 原生 Python", "实现深度正射投影、GraphCut、owner 与 MultiBand 管线。", "panorama.jpg、报告、delivery.json"),
    ], [1450, 2200, 3450, 2260])

    h2(doc, "数据与坐标约定")
    add_bullets(doc, [
        "采集会话以 RGB 彩色坐标系为基准；深度图必须是对齐后的 depth_aligned 图像。",
        "项目内部距离、平移、RMSE 和表面深度一律采用毫米；只有 RTAB-Map/Open3D 适配边界可临时转为米。",
        "camera_to_world 是 4×4 齐次 SE(3)，将彩色相机坐标映射到首个可信节点的世界坐标。",
        "二维侧扫画布由 scan_axis、up_axis、normal_axis 构成：横向为移动方向，纵向为相机上方向，法向只用于可见面 z-buffer。",
        "每一个 RTAB-Map 输出位姿必须能映射回唯一原始 RGB-D frame_id；无法关联的节点不能作为正式渲染源。",
    ])

    h1(doc, "四、RTAB-Map 三维建图设计")
    h2(doc, "Windows 原生采集与 RTAB-Map 桥接")
    add_text(doc, "采集阶段包含两种无界面模式。flash_video_capture_adapter 在 Windows 原生进程中启动 Gemini 305 的连续视频 Pipeline，复用 Flash 的设备枚举、共同 RGB-D profile 选择、D2C 对齐、帧同步和基础属性读取能力，并持续取得同步帧。每个连续视频帧按原始顺序写入彩色图、对齐深度和元数据。")
    add_text(doc, "flash_photo_capture_adapter 复用 Flash 的软件触发与同步取帧能力，以 CLI/API 请求单次 RGB-D 照片。每次请求只发起一次正式触发，并等待该触发产生的新鲜完整同步帧；触发、收帧、解码或写盘失败时不自动重试触发。照片模式不包含桌面入口，但可在相机移动至离散拍摄位置后多次调用，所有照片追加写入同一个严格 RGB-D 会话。")
    add_text(doc, "建图阶段由 rtabmap_native_bridge.exe 在 Windows 上读取会话目录并调用 RTAB-Map C++ 库。桥接程序把每个已验证的 RGB-D 帧、彩色相机模型、时间戳和深度比例转换为 RTAB-Map 传感器数据，运行 RGB-D 里程计与增量图优化，持久化 map.db，并导出关键帧位姿、点云/可选网格和结构化审计结果。")
    add_table(doc, ["节点/组件", "订阅或输入", "输出", "责任"], [
        ("flash_video_capture_adapter", "Gemini 305 原生连续视频流", "同步 RGB、aligned depth、内参、时间戳、曝光", "无界面复用 Flash 的设备控制能力，持续输出合格 RGB-D 帧。"),
        ("flash_photo_capture_adapter", "CLI/API 单次采集请求", "新鲜同步 RGB-D 帧、触发序号、时间戳、曝光、失败状态", "无界面复用 Flash 的软件触发与同步取帧；一次请求只允许一次正式触发。"),
        ("rtabmap_native_bridge.exe", "会话目录中的 RGB-D 与标定", "相对位姿、图、map.db、优化节点", "Windows C++ RTAB-Map 调用、连续 RGB-D 里程计与图优化。"),
        ("rtabmap_export_adapter", "map.db 与桥接导出", "关键帧位姿、点云/网格、审计 JSON", "转换坐标与单位，关联原始帧，执行正式质量校验。"),
    ], [1800, 2550, 2300, 2710])

    h2(doc, "RTAB-Map 参数与审计原则")
    add_bullets(doc, [
        "启用 RGB-D 输入和 RGB-D 视觉里程计；对每个节点记录时间戳、原始 frame_id、内参标识与深度比例。",
        "GPS 不进入本项目的室内正式链路；IMU 不是必需输入。未来如加入 IMU，只可作为姿态/里程计增强，不能替代 RGB-D 几何验证。",
        "回环候选不得直接成为正式结果。必须导出并检查其几何一致性、SE(3) 有限性、残差及对单向侧扫轨迹的影响。",
        "若 RTAB-Map 图边不满足当前工程“仅接受审计后的 RGB-D 几何约束”的要求，应在适配器中拒绝或降为诊断信息，而非放宽正式门限。",
        "三维点云和网格用于人工核验、档案和后续三维应用；二维全景仍由原始 RGB-D 重投影生成。",
    ])

    h2(doc, "三维成果生成")
    add_text(doc, "建图结束后保留 RTAB-Map 数据库作为可复算证据，并导出带颜色的稠密点云。网格为可选成果：只有点云密度、视角覆盖和深度质量满足要求时才生成；网格缺洞不能通过平滑或补面伪造。三维导出报告必须包含节点数、边数、回环数、有效深度比例、位姿图连通性、尺度单位和导出版本。")

    h1(doc, "五、三维到二维侧扫全景设计")
    h2(doc, "为何不直接展开网格")
    add_text(doc, "连续平移采集和离散照片采集都会产生真实视差。将网格纹理直接做 UV 展开或从固定视点截图，会丢失原始相机可见性和遮挡关系，近景物体极易出现拉伸、重影与接缝错位。系统应采用基于原始 RGB-D 的点投影与 z-buffer，并将 RTAB-Map 的价值限定为优化位姿与三维验证，而不是替换渲染几何。")

    h2(doc, "正射投影算法")
    add_numbers(doc, [
        "从通过审计的 RTAB-Map 关键帧中选取最多 32 个渲染源，要求覆盖率不少于 95%，相邻投影足迹重叠不少于 34%。",
        "由关键帧相机中心与方向估计统一扫描轴、上轴和法向；拒绝方向不一致、明显反向或上下/前后漂移过大的轨迹。",
        "对每帧原始全分辨率 RGB-D 像素执行深度回投影、camera_to_world 变换和正射坐标映射。",
        "每源使用独立 z-buffer，输出 warped_rgb、valid_mask、surface_depth_mm、camera_depth_mm 与投影统计；不跨深度断层三角化，不补造空洞。",
        "在统一画布中进行相邻帧接缝竞争与融合，再生成最大有效矩形裁剪后的 panorama.jpg。",
    ])

    h2(doc, "接缝与融合策略")
    add_table(doc, ["阶段", "输入", "规则", "失败条件"], [
        ("风险识别", "共同有效区、颜色、梯度、世界深度、相机深度", "识别近景、小于 1 m 区域、遮挡和深度不连续风险。", "没有真实共同有效区。"),
        ("硬 owner", "高风险连通域", "每一风险区域仅指定一个可靠来源，禁止双删或双 owner。", "不存在唯一可靠 owner 或安全通道被切断。"),
        ("GraphCut", "相邻 pair corridor", "只在安全走廊调用颜色/梯度 GraphCut；非相邻帧不竞争。", "GraphCut 异常、边界缺失或 owner 拓扑不成立。"),
        ("MultiBand", "通过审计的 owner 边界窄带", "每个相邻边界独立局部融合，窄带外直接复制唯一 owner。", "输出 mask 不完整、风险带被过度融合或出现黑洞。"),
    ], [1400, 2450, 3500, 2010])

    h1(doc, "六、软件改造设计")
    add_table(doc, ["新增/修改对象", "改造内容", "完成判据"], [
        ("configs/rtabmap_windows.yaml", "新增 Windows 原生 RTAB-Map 库/桥接程序路径、视频/照片会话参数、正式门限及导出目录。正式配置只能收紧，不得放宽质量包络。", "配置可加载；错误路径可诊断；正式模式不接受不安全覆盖。"),
        ("src/panorama_demo/flash_video_capture_adapter.py", "封装 Flash 可复用的设备枚举、profile 选择、D2C 对齐、帧同步和连续视频 Pipeline。", "连续 RGB-D 流可稳定保存；采集逻辑不依赖交互界面。"),
        ("src/panorama_demo/flash_photo_capture_adapter.py", "封装 Flash 可复用的软件触发、同步取帧和单帧写盘能力，以 CLI/API 实现无界面照片采集。", "单次请求对应一次正式触发；失败不自动重触发。"),
        ("native/rtabmap_native_bridge", "新增 Windows C++ 可执行程序，读取会话、调用 RTAB-Map 库、保存 map.db、导出位姿/点云并归一化错误。", "失败不留下 delivery.json；返回结构化错误。"),
        ("src/panorama_demo/rtabmap_export.py", "读取 Windows 桥接导出，统一单位、坐标和优化位姿；输出三维资产与审计 JSON。", "位姿为有限正交 SE(3)，图连通且单位正确。"),
        ("stitch_sequence.py", "实现 pose_backend 编排，串联 RTAB-Map 适配器、投影、渲染、门禁和原子交付。", "RTAB-Map 路径不导入旧 UniStitch/Torch 回退。"),
        ("tests/", "增加连续视频与无界面照片采集、帧关联、单位、回环审计、导出失败、RTAB-Map 到渲染集成测试。", "合成回归、错误路径、真实会话验收分别可报告。"),
    ], [2150, 4550, 2660])

    h2(doc, "命令行建议")
    add_text(doc, "正式用户入口保持简单，不暴露算法调参：")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("g305-capture --output .\\data\\captures\n"
                  "g305-panorama .\\data\\captures\\run_YYYYMMDD_HHMMSS --backend rtabmap_rgbd --output .\\outputs\\sequence")
    set_run_font(r, name="Consolas", size=9.5, color=INK)
    add_text(doc, "内部适配器直接启动 Windows 原生 rtabmap_native_bridge.exe；用户只提供会话目录与输出目录。Flash 复用组件、OrbbecSDK、RTAB-Map 库、桥接程序和 Python 依赖均应锁定版本并写入 Windows 安装脚本，避免环境漂移。")

    h1(doc, "七、质量门禁、异常处理与交付")
    h2(doc, "正式质量门禁")
    add_table(doc, ["类别", "硬性要求", "处理策略"], [
        ("输入结构", "必须有有效标定、aligned depth、正深度比例、时间戳、曝光与一致尺寸。", "结构失败立即拒绝；诊断模式也不得绕过。"),
        ("位姿", "图连通、SE(3) 有限且正交、相邻必要边通过、无明显逆向/垂直/前后漂移。", "拒绝生成三维正式资产和全景交付。"),
        ("投影", "画布和聚合工作集不超过 200 MP；最多 32 个源；深度空洞不得伪造。", "预分配前失败，避免内存耗尽或部分输出。"),
        ("接缝", "严格单 owner、非相邻 owner 不接触、风险 <= 0.10、无黑边/缺口。", "拒绝发布；输出 failure.json。"),
        ("交付", "正式文件先写 pending，再 os.replace；delivery.json 最后写入且 quality_pass=true。", "任何普通异常清理正式/诊断残留并记录失败。"),
    ], [1450, 5100, 2810])

    h2(doc, "验收标准")
    add_bullets(doc, [
        "软件验收：采集、会话、位姿、投影、渲染、交付与 RTAB-Map 适配层的单元、合成、错误路径与集成测试全部通过。",
        "三维验收：可生成可打开的 map.db 与彩色点云；报告中的节点、边、位姿和单位可追溯。",
        "二维验收：最近约 0.5 m 物体无明显重影和前景拉伸；有效区域无黑边、无接缝孔洞；交付目录只在通过时含 delivery.json。",
        "现场验收：连续视频模式在静止、0.5、1.0、1.5 m/s 下采集；照片模式在多个离散位置连续触发，确认 queue_drops=0、write_errors=0、RGB-D 同步、曝光合规和质量门禁结果。",
        "可维护性验收：采集、建图、渲染和交付均可在同一 Windows 主机完成；可从固定会话和固定 Windows 依赖版本复现建图结果。",
    ])

    h1(doc, "八、实施计划与里程碑")
    add_table(doc, ["阶段", "周期", "工作内容", "里程碑产物"], [
        ("P0 基线冻结", "第 1 周", "锁定 Windows 会话格式、视频/照片 profile、质量门禁和依赖版本；建立 Windows 原生 RTAB-Map 编译/运行环境。", "环境说明、版本锁定、基线测试报告。"),
        ("P1 双模式采集", "第 2 周", "实现 Flash 功能复用适配层，验证连续视频流、无界面软件触发、D2C 对齐、帧同步、frame_id/stamp 映射和写盘。", "视频与照片会话、输入审计。"),
        ("P2 三维建图", "第 3 周", "实现 Windows C++ RTAB-Map 桥接，接入 RGB-D odometry；导出数据库、点云、关键帧位姿和审计数据。", "三维成果与 rtabmap_transforms.json。"),
        ("P3 二维联调", "第 4 周", "把 RTAB-Map 位姿接入现有投影、GraphCut、owner 与 MultiBand 管线。", "端到端合成全景与失败清理测试。"),
        ("P4 实机验收", "第 5 周", "Gemini 305 静止与多速度现场测试；调整采集照明和操作流程，不放宽正式门限。", "现场采样报告与问题清单。"),
        ("P5 固化交付", "第 6 周", "完成性能、文档、回归、部署脚本和交付验收。", "发布包、操作手册、验收报告。"),
    ], [1350, 900, 4570, 2540])

    h1(doc, "九、风险与控制措施")
    add_table(doc, ["风险", "影响", "控制措施"], [
        ("Flash 功能与项目采集器的接口差异", "视频流无法稳定启动、单次触发行为不唯一或元数据不完整。", "仅抽取已验证的设备控制能力；建立独立视频/照片适配层及 profile/时间戳/对齐/触发回读测试，隔离与采集无关的功能。"),
        ("长曝光与快速移动", "视觉里程计失效、轨迹漂移和重影。", "保持 800 µs 采集目标与 1200 µs 正式上限；增加补光而非提高门限。"),
        ("低纹理、反光、动态物体", "RGB-D 里程计、回环和接缝不可靠。", "输入质量门禁、现场规程、明确失败；不使用 2D 回退或图像平均掩盖问题。"),
        ("RTAB-Map 回环误约束", "三维图和二维投影整体变形。", "导出后审计几何约束、轨迹方向和残差；不合格回环不进入正式结果。"),
        ("内存与画布过大", "渲染崩溃或系统失稳。", "200 MP 画布/聚合工作集硬限、32 源上限和预分配检查。"),
        ("环境版本漂移", "复现实验与部署失败。", "锁定 Windows SDK、Flash 复用组件、RTAB-Map 库、桥接程序、Python 依赖和配置文件。"),
    ], [2100, 2700, 4570])

    h1(doc, "十、资源需求与运维")
    add_table(doc, ["资源", "建议规格", "用途"], [
        ("采集主机", "Windows 10/11、USB 3、NVMe SSD、16 GB 以上内存", "驱动 Gemini 305、稳定写入 RGB-D 会话、运行二维渲染。"),
        ("建图环境", "Windows 原生 RTAB-Map C++ 运行时；16 GB 以上内存，建议独立 SSD 空间", "RGB-D 会话读取、图优化、数据库与点云/网格导出。"),
        ("相机与照明", "Gemini 305、稳定支架/滑轨、连续补光", "减少长曝光、运动模糊、深度噪声与反光。"),
        ("数据存储", "按会话保留原始 RGB-D、map.db、报告与最终输出", "实现可追溯、可复算与问题定位。"),
    ], [1850, 3800, 3720])

    h1(doc, "十一、结论与立项建议")
    add_text(doc, "本项目将建设一套完整的 Windows RGB-D 采集、三维建图与二维侧扫全景系统，并以严格输入约束、深度几何投影、接缝防护、原子交付和自动化验收作为统一质量基础。RTAB-Map 采用离线、可审计方式承担位姿图优化和三维成果生成；二维全景由原始 RGB-D 的正射重投影与深度约束融合完成。")
    add_text(doc, "建议按六周计划立项实施：先完成 Flash 功能复用的视频与无界面照片采集及位姿关联，再完成 Windows 原生 RTAB-Map 三维导出和二维渲染联调，最后进行实机多速度验收。该路径全程在 Windows 中运行，避免将三维网格简单二维化导致的近景视差和遮挡失真，并能保留完整的三维资产与质量证据。")
    add_callout(doc, "立项结论", "建议立项。建设范围明确，所有模块均以统一的 fail-closed 质量门禁和可追溯交付为目标开展实施。")

    h1(doc, "附录 A：参考依据")
    add_bullets(doc, [
        "RTAB-Map 官方项目主页：RGB-D、Stereo 与 Lidar 图优化 SLAM；https://introlab.github.io/rtabmap/",
        "RTAB-Map Kinect mapping 指南：视觉里程计对纹理、深度范围、运动模糊及 RGB/Depth 配准的要求；https://github.com/introlab/rtabmap/wiki/Kinect-mapping",
    ])

    doc.core_properties.title = "Gemini 305 RGB-D 三维建图与移动侧扫全景系统项目实施书"
    doc.core_properties.subject = "Windows + RTAB-Map 实施方案"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated from the current project state on 2026-07-13."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
